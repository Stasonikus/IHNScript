import os
import json
from docx import Document
from datetime import datetime

# ==========================================
# 🔥 Универсальное определение базовой папки
# ==========================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))      # путь к /utils
PROJECT_DIR = os.path.dirname(BASE_DIR)                    # путь к корню проекта

# Папки проекта
TEMPLATE_DIR = os.path.join(PROJECT_DIR, "templates")
DATA_DIR = os.path.join(PROJECT_DIR, "data")
OUTPUT_DIR = os.path.join(PROJECT_DIR, "output")

# Файлы
CARS_FILE = os.path.join(DATA_DIR, "cars.json")
COUNTER_FILE = os.path.join(DATA_DIR, "contract_counter.json")

# Автосоздание директорий
os.makedirs(TEMPLATE_DIR, exist_ok=True)
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)


# =======================================================
# 🔥 Функция поиска файла по всему проекту (на всякий случай)
# =======================================================
def find_file(filename, search_dir):
    for root, dirs, files in os.walk(search_dir):
        if filename in files:
            return os.path.join(root, filename)
    return None


# Если файл не найден — пробуем найти в любом месте проекта
if not os.path.exists(CARS_FILE):
    alt = find_file("cars.json", PROJECT_DIR)
    if alt:
        CARS_FILE = alt


# ==========================================
# 🔧 Функция замены текста в docx
# ==========================================
def replace_text_in_doc(doc, replacements):
    for paragraph in doc.paragraphs:
        full_text = ''.join(run.text for run in paragraph.runs)
        new_text = full_text
        for key, value in replacements.items():
            new_text = new_text.replace(key, str(value))
        if new_text != full_text:
            for i in range(len(paragraph.runs) - 1, -1, -1):
                p = paragraph.runs[i]._element
                p.getparent().remove(p)
            paragraph.add_run(new_text)

    # Таблички
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = ''.join(run.text for run in paragraph.runs)
                    new_text = full_text
                    for key, value in replacements.items():
                        new_text = new_text.replace(key, str(value))
                    if new_text != full_text:
                        for i in range(len(paragraph.runs) - 1, -1, -1):
                            p = paragraph.runs[i]._element
                            p.getparent().remove(p)
                        paragraph.add_run(new_text)


# ==========================================
# 🔢 Контрактный номер
# ==========================================
def load_contract_number():
    if not os.path.exists(COUNTER_FILE):
        return 1
    with open(COUNTER_FILE, "r", encoding="utf-8") as f:
        data = json.load(f)
    return data.get("last_number", 0) + 1


def save_contract_number(number):
    with open(COUNTER_FILE, "w", encoding="utf-8") as f:
        json.dump({"last_number": number}, f, ensure_ascii=False, indent=4)


# ==========================================
# 📝 Генерация документов
# ==========================================
def generate_docs(data_dict, client_name):
    if not os.path.exists(TEMPLATE_DIR):
        print(f"❌ Папка шаблонов не найдена: {TEMPLATE_DIR}")
        return

    for filename in os.listdir(TEMPLATE_DIR):
        if filename.endswith(".docx"):
            template_path = os.path.join(TEMPLATE_DIR, filename)

            if "contract" in filename.lower():
                new_name = f"contract_{client_name}.docx"
            elif "poa" in filename.lower():
                new_name = f"poa_{client_name}.docx"
            elif "waybill" in filename.lower():
                new_name = f"waybill_{client_name}.docx"
            else:
                new_name = f"{client_name}_{filename}"

            output_path = os.path.join(OUTPUT_DIR, new_name)
            print(f"📄 Обрабатываю шаблон: {filename}")

            doc = Document(template_path)
            replace_text_in_doc(doc, data_dict)
            doc.save(output_path)

            print(f"✅ Файл сохранён: {output_path}\n")


# ==========================================
# 🚗 Загрузка списка машин
# ==========================================
def load_cars():
    if not os.path.exists(CARS_FILE):
        print(f"❌ Не найден cars.json: {CARS_FILE}")
        return []
    with open(CARS_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


# ==========================================
# 📆 Подсчёт дней
# ==========================================
def calculate_days(start_date, end_date):
    fmt = "%d.%m.%Y"
    d1 = datetime.strptime(start_date, fmt)
    d2 = datetime.strptime(end_date, fmt)
    return (d2 - d1).days + 1


# ==========================================
# 🛣️ Выбор покрытия дорог
# ==========================================
def choose_road_types():
    road_options = ["Paved", "Gravel", "Dirt Tracks", "Off-Road", "Asphalt"]
    print("\n🛣️ Доступные типы дорог:")
    for i, r in enumerate(road_options, 1):
        print(f"{i}. {r}")
    print("\nМожно выбрать несколько (1,3,5)")
    choice = input("Выберите типы дорог: ").strip()
    selected = []
    if choice:
        for ch in choice.split(","):
            if ch.strip().isdigit():
                idx = int(ch.strip())
                if 1 <= idx <= len(road_options):
                    selected.append(road_options[idx - 1])
    return ", ".join(selected) if selected else "Asphalt"


# ==========================================
# 🌍 Новый выбор дополнительных стран
# Казахстан — всегда по умолчанию!
# ==========================================

def choose_additional_countries():
    """
    Менеджер выбирает только доп. страны.
    Казахстан всегда добавляется автоматически.
    """
    options = {
        1: ("Kyrgyzstan", "Кыргызстан"),
        2: ("Uzbekistan", "Узбекистан"),
        3: ("Tajikistan", "Таджикистан")
    }

    print("\n🌍 Выберите дополнительные страны (кроме Казахстана):")
    for i, (en, ru) in options.items():
        print(f"{i}. {ru}")

    print("\nМожно выбрать несколько (например: 1,2) или оставить пусто")
    choice = input("Ваш выбор: ").strip()

    selected_eng = []   # Для контракта (EN)
    selected_rus = []   # Для доверенности (RU)

    if choice:
        for ch in choice.split(","):
            ch = ch.strip()
            if ch.isdigit() and int(ch) in options:
                en, ru = options[int(ch)]
                selected_eng.append(en)
                selected_rus.append(ru)

    return selected_eng, selected_rus


# ==========================================
# 🌍 Формирование финальных строк для документов
# ==========================================

def format_country_strings(selected_eng, selected_rus):
    """
    Возвращает 2 строки:
    1) Для {{ALLOWED_COUNTRIES}} — всегда начинается с Kazakhstan
    2) Для {{ALLOWED_TERRITORIES}} — только русские страны (для POA & Waybill)
    """

    # Казахстан всегда первая страна
    countries_for_contract = ["Kazakhstan"] + selected_eng

    # Для доверенности / waybill Казахстан НЕ включается
    territories_ru = ", ".join(selected_rus) if selected_rus else ""

    return ", ".join(countries_for_contract), territories_ru



# ==========================================
# 🚀 Основной запуск
# ==========================================
if __name__ == "__main__":
    cars = load_cars()
    if not cars:
        print("🚫 Нет данных об автомобилях! Добавьте cars.json в папку /data")
        exit()

    print("🚗 Доступные автомобили:")
    for i, car in enumerate(cars, 1):
        print(f"{i}. {car['make']} {car['model']} ({car['plate']})")

    choice = int(input("Выбери номер машины: ")) - 1
    selected_car = cars[choice]

    print("\nВведите данные клиента:")
    client_name = input("Имя клиента (Фамилия Имя): ")
    date_of_birth = input("Дата рождения (ДД.ММ.ГГГГ): ")
    address = input("Адрес проживания: ")
    phone = input("Телефон: ")
    email = input("Email: ")

    passport_number = input("Номер паспорта: ")
    passport_issue_date = input("Дата выдачи паспорта: ")
    passport_issue_by = input("Кем выдан паспорт: ")
    license_num = input("Номер ВУ: ")

    start_date = input("Дата начала аренды: ")
    end_date = input("Дата конца аренды: ")

    rental_rate = float(input("Цена за сутки (USD): "))
    days = calculate_days(start_date, end_date)
    total_amount = rental_rate * days
    security_deposit = float(input("Сумма залога (USD): "))

    # Доп. водители
    print("\nЕсть ли дополнительные водители? (да/нет)")
    add_drivers = input().strip().lower()
    driver_data = {
        "{{DRIVER1_NAME}}": "",
        "{{DRIVER1_LICENSE}}": "",
        "{{DRIVER2_NAME}}": "",
        "{{DRIVER2_LICENSE}}": "",
        "{{DRIVER3_NAME}}": "",
        "{{DRIVER3_LICENSE}}": ""
    }

    if add_drivers == "да":
        num = int(input("Сколько водителей (до 3)? "))
        for i in range(num):
            name = input(f"Имя водителя {i+1}: ")
            lic = input(f"Права водителя {i+1}: ")
            driver_data[f"{{{{DRIVER{i+1}_NAME}}}}"] = name
            driver_data[f"{{{{DRIVER{i+1}_LICENSE}}}}"] = lic

    road_types = choose_road_types()

    # Контракт
    contract_number = load_contract_number()
    contract_date = datetime.now().strftime("%d.%m.%Y")
    save_contract_number(contract_number)

    # Выбор регионов
    # Новый выбор стран
    selected_eng, selected_rus = choose_additional_countries()

    allowed_countries, allowed_territories_ru = format_country_strings(
        selected_eng,
        selected_rus
    )


    # Данные
    data = {
        "{{CONTRACT_DATE}}": contract_date,
        "{{CONTRACT_NUMBER}}": contract_number,
        "{{CLIENT_NAME}}": client_name,
        "{{DATE_OF_BIRTH}}": date_of_birth,
        "{{ADDRESS}}": address,
        "{{PHONE}}": phone,
        "{{EMAIL}}": email,
        "{{PASSPORT_NUMBER}}": passport_number,
        "{{PASSPORT_ISSUE_DATE}}": passport_issue_date,
        "{{PASSPORT_ISSUE_BY}}": passport_issue_by,
        "{{DRIVER_LICENSE}}": license_num,
        "{{RENTAL_START}}": start_date,
        "{{RENTAL_END}}": end_date,
        "{{RENTAL_RATE}}": rental_rate,
        "{{TOTAL_AMOUNT}}": f"{total_amount:.2f}",
        "{{SECURITY_DEPOSIT}}": f"{security_deposit:.2f}",

        # Машина
        "{{CAR_MAKE}}": selected_car["make"],
        "{{CAR_MODEL}}": selected_car["model"],
        "{{CAR_NAME}}": f"{selected_car['make']} {selected_car['model']}",
        "{{CAR_YEAR}}": selected_car["year"],
        "{{CAR_COLOR}}": selected_car["color"],
        "{{CAR_PLATE}}": selected_car["plate"],
        "{{CAR_VIN}}": selected_car["vin"],

        # Территории
        "{{ALLOWED_COUNTRIES}}": allowed_countries,
        "{{ALLOWED_TERRITORIES}}": allowed_territories_ru,
        "{{TERRITORIES_FOR_POA}}": f"по всей территории Казахстана" + 
                                  (f" и за её пределами: {allowed_territories_ru}" if allowed_territories_ru else "")

    }

    data.update(driver_data)

    generate_docs(data, client_name.replace(" ", "_"))
